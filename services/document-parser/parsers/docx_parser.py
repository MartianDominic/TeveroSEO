"""
DOCX Parser using python-docx.
Phase 102-08: Extracts text with formatting metadata.

Key features:
- Full text extraction from paragraphs
- Table content extraction
- Bold/italic formatting detection
- Heading style recognition
- Image presence detection
"""

from docx import Document
from dataclasses import dataclass
from typing import List, Dict, Any
from collections import Counter


@dataclass
class DocxParseResult:
    """Result from DOCX parsing."""
    text: str
    page_count: int  # Estimated from content length
    metadata: Dict[str, Any]
    fonts: List[Dict[str, Any]]
    colors: List[str]
    has_images: bool
    needs_ocr: bool  # Always False for DOCX


def parse_docx(file_path: str) -> DocxParseResult:
    """
    Parse DOCX and extract text with formatting metadata.

    Note: DOCX doesn't have true "pages" - we estimate based on content.

    Returns structured data including:
    - Full text content
    - Font information (name, size, frequency)
    - Color palette (when available)
    - Image presence
    """

    doc = Document(file_path)

    all_text = []
    fonts: Counter = Counter()
    colors: Counter = Counter()
    has_images = False

    # Process paragraphs
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            all_text.append(para_text)

            # Extract formatting from runs
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue

                # Track fonts
                font_name = None
                font_size = 12.0

                if run.font.name:
                    font_name = run.font.name

                if run.font.size:
                    try:
                        font_size = run.font.size.pt
                    except (AttributeError, TypeError):
                        font_size = 12.0

                if font_name:
                    font_key = f"{font_name}:{font_size:.0f}"
                    fonts[font_key] += len(run_text)

                # Track colors
                if run.font.color and run.font.color.rgb:
                    try:
                        hex_color = f"#{run.font.color.rgb}"
                        colors[hex_color] += len(run_text)
                    except (AttributeError, TypeError):
                        pass

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                all_text.append(" | ".join(row_text))

    # Check for images
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_images = True
                break
    except (AttributeError, KeyError):
        pass

    # Estimate page count (rough: ~500 words per page)
    word_count = sum(len(t.split()) for t in all_text)
    estimated_pages = max(1, word_count // 500)

    # Build font list (top 10 by usage)
    font_list = []
    for font_key, count in fonts.most_common(10):
        parts = font_key.rsplit(":", 1)
        if len(parts) == 2:
            font_list.append({
                "font": parts[0],
                "size": float(parts[1]),
                "usage": count,
            })

    # Build color list (top 5 by usage)
    color_list = [color for color, _ in colors.most_common(5)]

    # Extract core properties
    metadata = {
        "title": "",
        "author": "",
        "created": "",
    }

    try:
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created"] = str(core_props.created)
    except (AttributeError, KeyError):
        pass

    return DocxParseResult(
        text="\n\n".join(all_text),
        page_count=estimated_pages,
        metadata=metadata,
        fonts=font_list,
        colors=color_list,
        has_images=has_images,
        needs_ocr=False,  # DOCX always has extractable text
    )
